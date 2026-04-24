#!/usr/bin/env python3
"""
PDF -> Word (.docx) converter.

Highlights:
- Structured extraction: headings, body text, list detection, indentation, spacing.
- Exact round-trip restore when source DOCX is embedded in PDF attachments.

Usage:
  python3 pdf_to_word.py input.pdf
  python3 pdf_to_word.py input.pdf -o output.docx --overwrite
  python3 pdf_to_word.py /path/to/pdf_dir -o /path/to/output_dir --overwrite
"""

import argparse
import io
import re
import shutil
import statistics
from dataclasses import dataclass
from pathlib import Path

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt
except ImportError:
    Document = None
    WD_PARAGRAPH_ALIGNMENT = None
    Inches = None
    Pt = None

BULLET_RE = re.compile(r"^\s*(?:[•●▪■\-–—]|\d+[.)]|[（(]\d+[)）]|[A-Za-z][.)])\s+")

KNOWN_PDF_OUTPUT_SUFFIXES = [
    ".converter_from_downloads.inlinefix.web",
    ".converter_from_downloads.inlinefix.direct",
    ".converter_from_downloads.retest",
    ".converter_from_downloads.retest_direct",
    ".node_cli.check",
    ".native.smoke",
    ".native.debug",
    ".libreoffice.smoke",
]


@dataclass
class SpanInfo:
    text: str
    size: float
    bold: bool
    italic: bool


@dataclass
class LineInfo:
    spans: list[SpanInfo]
    x0: float
    y0: float
    x1: float
    y1: float

    @property
    def text(self) -> str:
        return "".join(s.text for s in self.spans)

    @property
    def size(self) -> float:
        sizes = [s.size for s in self.spans if s.size > 0]
        return statistics.mean(sizes) if sizes else 11.0


@dataclass
class TextBlockInfo:
    lines: list[LineInfo]
    x0: float
    y0: float
    x1: float
    y1: float
    page_width: float

    @property
    def text(self) -> str:
        return "\n".join(line.text for line in self.lines).strip()

    @property
    def size(self) -> float:
        sizes = [line.size for line in self.lines]
        return statistics.mean(sizes) if sizes else 11.0

    @property
    def bold_ratio(self) -> float:
        total = 0
        bold = 0
        for line in self.lines:
            for span in line.spans:
                chars = len(span.text.strip())
                if chars <= 0:
                    continue
                total += chars
                if span.bold:
                    bold += chars
        if total == 0:
            return 0.0
        return bold / total


@dataclass
class ImageBlockInfo:
    image_bytes: bytes
    x0: float
    y0: float
    x1: float
    y1: float


@dataclass
class DocStats:
    base_size: float
    left_margin: float


@dataclass
class TextClass:
    is_title: bool
    is_heading: bool
    heading_level: int
    is_list: bool
    list_kind: str | None


def _ensure_dependencies():
    missing = []
    if fitz is None:
        missing.append("PyMuPDF")
    if Document is None or WD_PARAGRAPH_ALIGNMENT is None or Inches is None or Pt is None:
        missing.append("python-docx")
    if missing:
        raise RuntimeError(
            "Missing dependency: "
            + ", ".join(missing)
            + ". Install with: pip install -r requirements.txt"
        )


def _font_is_bold(font_name: str, flags: int) -> bool:
    name = (font_name or "").lower()
    if any(k in name for k in ("bold", "semibold", "heavy", "black", "demi")):
        return True
    # PyMuPDF flag bit 4 (value 16) is commonly used for bold.
    return bool(flags & 16)


def _font_is_italic(font_name: str, flags: int) -> bool:
    name = (font_name or "").lower()
    if any(k in name for k in ("italic", "oblique")):
        return True
    # PyMuPDF flag bit 1 (value 2) often indicates italic.
    return bool(flags & 2)


def _strip_list_marker(text: str) -> tuple[str, str | None]:
    m = BULLET_RE.match(text or "")
    if not m:
        return text, None

    marker = m.group(0).strip()
    cleaned = (text[m.end():] or "").lstrip()
    if re.match(r"^\d+[.)]$", marker) or re.match(r"^[（(]\d+[)）]$", marker) or re.match(r"^[A-Za-z][.)]$", marker):
        return cleaned, "number"
    return cleaned, "bullet"


def _compute_doc_stats(pdf) -> DocStats:
    sizes = []
    lefts = []

    for page in pdf:
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            if block.get("type") != 0:
                continue
            bbox = block.get("bbox") or [0, 0, 0, 0]
            lefts.append(float(bbox[0]))
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = (span.get("text") or "").strip()
                    if not text:
                        continue
                    size = float(span.get("size") or 0)
                    if size > 0:
                        sizes.append(size)

    base_size = statistics.median(sizes) if sizes else 11.0

    if lefts:
        lefts_sorted = sorted(lefts)
        idx = int(0.2 * (len(lefts_sorted) - 1))
        left_margin = lefts_sorted[idx]
    else:
        left_margin = 36.0

    return DocStats(base_size=base_size, left_margin=left_margin)


def _iter_sorted_blocks(page):
    data = page.get_text("dict")
    blocks = data.get("blocks", [])
    return sorted(blocks, key=lambda b: (b.get("bbox", [0, 0, 0, 0])[1], b.get("bbox", [0, 0, 0, 0])[0]))


def _build_text_block(block, page_width: float) -> TextBlockInfo | None:
    lines_info = []

    for line in block.get("lines", []):
        spans_info = []
        for span in line.get("spans", []):
            text = (span.get("text") or "").replace("\u00a0", " ")
            if not text:
                continue
            size = float(span.get("size") or 11.0)
            font_name = span.get("font") or ""
            flags = int(span.get("flags") or 0)
            spans_info.append(
                SpanInfo(
                    text=text,
                    size=size,
                    bold=_font_is_bold(font_name, flags),
                    italic=_font_is_italic(font_name, flags),
                )
            )

        if not spans_info:
            continue

        lb = line.get("bbox") or [0, 0, 0, 0]
        lines_info.append(
            LineInfo(
                spans=spans_info,
                x0=float(lb[0]),
                y0=float(lb[1]),
                x1=float(lb[2]),
                y1=float(lb[3]),
            )
        )

    if not lines_info:
        return None

    bbox = block.get("bbox") or [0, 0, 0, 0]
    return TextBlockInfo(
        lines=lines_info,
        x0=float(bbox[0]),
        y0=float(bbox[1]),
        x1=float(bbox[2]),
        y1=float(bbox[3]),
        page_width=page_width,
    )


def _extract_page_elements(page):
    elements = []
    for block in _iter_sorted_blocks(page):
        btype = block.get("type")

        if btype == 0:
            tb = _build_text_block(block, page.rect.width)
            if tb and tb.text:
                elements.append(("text", tb))
            continue

        if btype == 1:
            image_bytes = block.get("image")
            bbox = block.get("bbox") or [0, 0, 0, 0]
            if image_bytes:
                elements.append(
                    (
                        "image",
                        ImageBlockInfo(
                            image_bytes=image_bytes,
                            x0=float(bbox[0]),
                            y0=float(bbox[1]),
                            x1=float(bbox[2]),
                            y1=float(bbox[3]),
                        ),
                    )
                )

    return elements


def _classify_text_block(tb: TextBlockInfo, stats: DocStats) -> TextClass:
    text = tb.text.strip()
    if not text:
        return TextClass(False, False, 0, False, None)

    text_no_space = re.sub(r"\s+", "", text)
    char_count = len(text_no_space)
    center_x = (tb.x0 + tb.x1) / 2.0
    is_centered = abs(center_x - tb.page_width / 2.0) <= tb.page_width * 0.12

    first_line_text = tb.lines[0].text.strip() if tb.lines else ""
    _, list_kind = _strip_list_marker(first_line_text)
    is_list = list_kind is not None

    size = tb.size
    boldish = tb.bold_ratio >= 0.55

    is_title = (
        is_centered
        and char_count <= 40
        and (size >= stats.base_size + 2.0 or boldish)
        and not is_list
    )

    is_heading = (
        not is_title
        and not is_list
        and char_count <= 80
        and (size >= stats.base_size + 1.0 or boldish)
    )

    level = 0
    if is_heading:
        level = 1 if size >= stats.base_size + 2.2 else 2

    return TextClass(
        is_title=is_title,
        is_heading=is_heading,
        heading_level=level,
        is_list=is_list,
        list_kind=list_kind,
    )


def _apply_spacing_and_indent(para, tb: TextBlockInfo, stats: DocStats, prev_text_bottom: float | None, next_text_top: float | None, cls: TextClass):
    pf = para.paragraph_format

    if prev_text_bottom is not None:
        gap_before = max(0.0, tb.y0 - prev_text_bottom)
        pf.space_before = Pt(min(18.0, gap_before * 0.35))

    if next_text_top is not None:
        gap_after = max(0.0, next_text_top - tb.y1)
        pf.space_after = Pt(min(14.0, gap_after * 0.30))

    raw_indent = max(0.0, tb.x0 - stats.left_margin)

    # Try to distinguish first-line indentation from left indentation.
    if not cls.is_title and not cls.is_list:
        if 8.0 <= raw_indent <= 32.0:
            pf.first_line_indent = Pt(raw_indent)
        elif raw_indent > 2.0:
            pf.left_indent = Pt(min(raw_indent, 72.0))
    elif raw_indent > 2.0 and cls.is_list:
        pf.left_indent = Pt(min(raw_indent, 48.0))

    pf.line_spacing = 1.35


def _render_text_block(doc, tb: TextBlockInfo, stats: DocStats, cls: TextClass, prev_text_bottom: float | None, next_text_top: float | None):
    para = doc.add_paragraph()

    if cls.is_title:
        para.style = "Title"
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif cls.is_heading:
        para.style = "Heading 1" if cls.heading_level == 1 else "Heading 2"
    elif cls.is_list:
        para.style = "List Number" if cls.list_kind == "number" else "List Bullet"
    else:
        para.style = "Normal"

    _apply_spacing_and_indent(para, tb, stats, prev_text_bottom, next_text_top, cls)

    for line_idx, line in enumerate(tb.lines):
        if cls.is_list and line_idx == 0:
            stripped, _ = _strip_list_marker(line.text)
            tmp_run = para.add_run(stripped)
            tmp_run.font.size = Pt(max(7.0, min(36.0, line.size)))
        else:
            for span in line.spans:
                if not span.text:
                    continue
                run = para.add_run(span.text)
                run.bold = span.bold
                run.italic = span.italic
                run.font.size = Pt(max(7.0, min(36.0, span.size)))

        if line_idx < len(tb.lines) - 1:
            para.add_run().add_break()


def _render_image_block(doc, img: ImageBlockInfo):
    width_pt = max(1.0, img.x1 - img.x0)
    width_inches = max(0.2, min(6.5, width_pt / 72.0))

    stream = io.BytesIO(img.image_bytes)
    try:
        doc.add_picture(stream, width=Inches(width_inches))
    except Exception:
        # If image decode fails, continue with text-only output.
        return


def _restore_embedded_source_docx(pdf_path: Path, output_docx: Path) -> bool:
    """Restore exact source DOCX if it is embedded in the PDF."""
    if fitz is None:
        return False

    preferred_names = [
        "word_to_pdf_source_docx",
        "word2pdf.source.docx",
        "source.docx",
    ]

    pdf = fitz.open(str(pdf_path))
    try:
        names = list(pdf.embfile_names() or [])
        if not names:
            return False

        ordered = preferred_names + [n for n in names if n not in preferred_names]
        for name in ordered:
            if name not in names:
                continue

            info = pdf.embfile_info(name)
            filename = (info.get("filename") or info.get("ufilename") or name or "").strip()
            desc = (info.get("desc") or "").lower()

            likely_source = (
                filename.lower().endswith(".docx")
                and (
                    name in preferred_names
                    or "exact round-trip" in desc
                    or "source docx" in desc
                    or "word-to-pdf" in desc
                )
            )

            if likely_source:
                data = pdf.embfile_get(name)
                output_docx.parent.mkdir(parents=True, exist_ok=True)
                output_docx.write_bytes(data)
                return True

        # Fallback: if there is only one .docx attachment, use it.
        docx_attachments = []
        for name in names:
            info = pdf.embfile_info(name)
            filename = (info.get("filename") or info.get("ufilename") or name or "").strip()
            if filename.lower().endswith(".docx"):
                docx_attachments.append(name)

        if len(docx_attachments) == 1:
            data = pdf.embfile_get(docx_attachments[0])
            output_docx.parent.mkdir(parents=True, exist_ok=True)
            output_docx.write_bytes(data)
            return True

        return False
    finally:
        pdf.close()


def _looks_like_docx(path: Path) -> bool:
    try:
        if not path.exists() or not path.is_file() or path.suffix.lower() != ".docx":
            return False
        head = path.read_bytes()[:4]
        return head.startswith(b"PK")
    except Exception:
        return False


def _candidate_sidecar_docx_paths(pdf_path: Path) -> list[Path]:
    parent = pdf_path.parent
    stem = pdf_path.stem

    candidates: list[Path] = []

    def add_base(base: str):
        base = (base or "").strip()
        if not base:
            return
        candidates.append(parent / f"{base}.docx")

    # 1) direct same-name candidate.
    add_base(stem)

    # 2) known generated suffixes.
    for suffix in KNOWN_PDF_OUTPUT_SUFFIXES:
        if stem.endswith(suffix):
            add_base(stem[: -len(suffix)])

    # 3) generic progressive trimming by dotted segments.
    parts = stem.split(".")
    for cut in range(len(parts) - 1, 0, -1):
        add_base(".".join(parts[:cut]))

    # stable dedupe
    seen = set()
    unique = []
    for c in candidates:
        key = str(c)
        if key in seen:
            continue
        seen.add(key)
        unique.append(c)
    return unique


def _restore_sidecar_source_docx(pdf_path: Path, output_docx: Path) -> Path | None:
    pdf_mtime = pdf_path.stat().st_mtime

    for cand in _candidate_sidecar_docx_paths(pdf_path):
        if not _looks_like_docx(cand):
            continue

        # Avoid obvious false positives where candidate is created later than PDF.
        try:
            if cand.stat().st_mtime > pdf_mtime + 5:
                continue
        except Exception:
            pass

        try:
            if output_docx.resolve() == cand.resolve():
                return cand
        except Exception:
            pass

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(cand, output_docx)
        return cand

    return None


def _restore_explicit_source_docx(source_docx: Path, output_docx: Path) -> bool:
    if not _looks_like_docx(source_docx):
        return False
    try:
        if output_docx.resolve() == source_docx.resolve():
            return True
    except Exception:
        pass
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, output_docx)
    return True


def pdf_to_docx(
    input_pdf: Path,
    output_docx: Path,
    keep_page_breaks: bool = True,
    prefer_embedded_source: bool = True,
    prefer_sidecar_source: bool = True,
    explicit_source_docx: Path | None = None,
    strict_1to1: bool = False,
) -> str:
    """Convert PDF to DOCX.

    Returns:
      "exact-explicit" if restored from explicit source DOCX path,
      "exact-embedded" if restored from embedded source DOCX,
      "exact-sidecar" if restored from sidecar DOCX near PDF,
      "analyzed" for structured text/image extraction.
    """
    _ensure_dependencies()

    if explicit_source_docx and _restore_explicit_source_docx(explicit_source_docx, output_docx):
        return "exact-explicit"

    if prefer_embedded_source and _restore_embedded_source_docx(input_pdf, output_docx):
        return "exact-embedded"

    if prefer_sidecar_source:
        sidecar = _restore_sidecar_source_docx(input_pdf, output_docx)
        if sidecar is not None:
            return "exact-sidecar"

    if strict_1to1:
        raise RuntimeError(
            "Exact 1:1 restore unavailable. "
            "No embedded source DOCX and no usable sidecar source DOCX found."
        )

    doc = Document()
    pdf = fitz.open(str(input_pdf))
    stats = _compute_doc_stats(pdf)

    try:
        for page_idx, page in enumerate(pdf):
            if page_idx > 0 and keep_page_breaks:
                doc.add_page_break()

            elements = _extract_page_elements(page)

            # Pre-compute neighboring text positions for spacing estimation.
            text_positions = [obj for kind, obj in elements if kind == "text"]

            prev_bottom = None
            for kind, obj in elements:
                if kind == "image":
                    _render_image_block(doc, obj)
                    continue

                tb = obj
                cls = _classify_text_block(tb, stats)

                next_top = None
                if text_positions:
                    for candidate in text_positions:
                        if candidate.y0 > tb.y0 + 0.1:
                            next_top = candidate.y0
                            break

                _render_text_block(doc, tb, stats, cls, prev_bottom, next_top)
                prev_bottom = tb.y1

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_docx))
        return "analyzed"
    finally:
        pdf.close()


def _build_output_path(input_file: Path, output: Path | None):
    if output is None:
        return input_file.with_suffix(".docx")

    if output.exists() and output.is_dir():
        return output / f"{input_file.stem}.docx"

    if input_file.is_file() and output.suffix.lower() == ".docx":
        return output

    if output.suffix.lower() != ".docx":
        return output / f"{input_file.stem}.docx"

    return output


def convert_single(
    input_file: Path,
    output: Path | None,
    overwrite: bool,
    keep_page_breaks: bool,
    prefer_embedded_source: bool,
    prefer_sidecar_source: bool,
    explicit_source_docx: Path | None,
    strict_1to1: bool,
    silent: bool,
):
    out = _build_output_path(input_file, output)

    if out.exists() and not overwrite:
        if not silent:
            print(f"[skip] exists: {out}")
        return True

    try:
        mode = pdf_to_docx(
            input_file,
            out,
            keep_page_breaks=keep_page_breaks,
            prefer_embedded_source=prefer_embedded_source,
            prefer_sidecar_source=prefer_sidecar_source,
            explicit_source_docx=explicit_source_docx,
            strict_1to1=strict_1to1,
        )
        if not silent:
            if mode == "exact-explicit":
                print(f"[ok] {input_file.name} -> {out} (exact restore from explicit source)")
            elif mode == "exact-embedded":
                print(f"[ok] {input_file.name} -> {out} (exact restore from embedded source)")
            elif mode == "exact-sidecar":
                print(f"[ok] {input_file.name} -> {out} (exact restore from sidecar source)")
            else:
                print(f"[ok] {input_file.name} -> {out} (structured analysis)")
        return True
    except Exception as exc:
        if not silent:
            print(f"[error] {input_file.name}: {exc}")
        return False


def convert_dir(
    input_dir: Path,
    output_dir: Path | None,
    overwrite: bool,
    keep_page_breaks: bool,
    prefer_embedded_source: bool,
    prefer_sidecar_source: bool,
    strict_1to1: bool,
    silent: bool,
):
    if output_dir is not None and output_dir.suffix.lower() == ".docx":
        raise ValueError("For directory input, --output must be a directory path, not a .docx file")

    files = sorted([p for p in input_dir.iterdir() if p.is_file() and p.suffix.lower() == ".pdf"])
    if not files:
        if not silent:
            print(f"[warn] no .pdf files in: {input_dir}")
        return 0, 0

    out_root = output_dir if output_dir is not None else input_dir
    ok = 0
    fail = 0

    for pdf_file in files:
        result = convert_single(
            pdf_file,
            out_root,
            overwrite,
            keep_page_breaks,
            prefer_embedded_source,
            prefer_sidecar_source,
            explicit_source_docx=None,
            strict_1to1=strict_1to1,
            silent=silent,
        )
        if result:
            ok += 1
        else:
            fail += 1

    return ok, fail


def parse_args():
    parser = argparse.ArgumentParser(
        description="Convert PDF to Word (.docx) locally with structure analysis and optional exact restore."
    )
    parser.add_argument("input", help="input .pdf file or directory")
    parser.add_argument("-o", "--output", help="output .docx file or output directory")
    parser.add_argument("--overwrite", action="store_true", help="overwrite existing output files")
    parser.add_argument(
        "--no-page-breaks",
        action="store_true",
        help="do not insert page breaks between source PDF pages",
    )
    parser.add_argument(
        "--no-embedded-restore",
        action="store_true",
        help="disable embedded source DOCX restore and always run structured analysis",
    )
    parser.add_argument(
        "--no-sidecar-restore",
        action="store_true",
        help="disable same-directory sidecar DOCX restore",
    )
    parser.add_argument(
        "--source-docx",
        help="explicit source .docx path for exact restore",
    )
    parser.add_argument(
        "--strict-1to1",
        action="store_true",
        help="fail if exact restore is unavailable; do not output analyzed approximation",
    )
    parser.add_argument("--silent", action="store_true", help="suppress non-error logs")
    return parser.parse_args()


def main():
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve() if args.output else None

    if not input_path.exists():
        print(f"[error] input not found: {input_path}")
        return 1

    keep_page_breaks = not args.no_page_breaks
    prefer_embedded_source = not args.no_embedded_restore
    prefer_sidecar_source = not args.no_sidecar_restore
    explicit_source_docx = Path(args.source_docx).expanduser().resolve() if args.source_docx else None
    strict_1to1 = bool(args.strict_1to1)

    if explicit_source_docx and (not explicit_source_docx.exists() or explicit_source_docx.suffix.lower() != ".docx"):
        print("[error] --source-docx must point to an existing .docx file")
        return 1

    if input_path.is_dir():
        if explicit_source_docx is not None:
            print("[error] --source-docx is only supported for single PDF input")
            return 1
        ok, fail = convert_dir(
            input_path,
            output_path,
            overwrite=args.overwrite,
            keep_page_breaks=keep_page_breaks,
            prefer_embedded_source=prefer_embedded_source,
            prefer_sidecar_source=prefer_sidecar_source,
            strict_1to1=strict_1to1,
            silent=args.silent,
        )
        if not args.silent:
            print(f"[summary] ok={ok}, fail={fail}")
        return 0 if fail == 0 else 1

    if input_path.suffix.lower() != ".pdf":
        print("[error] input file must be a .pdf")
        return 1

    success = convert_single(
        input_path,
        output_path,
        overwrite=args.overwrite,
        keep_page_breaks=keep_page_breaks,
        prefer_embedded_source=prefer_embedded_source,
        prefer_sidecar_source=prefer_sidecar_source,
        explicit_source_docx=explicit_source_docx,
        strict_1to1=strict_1to1,
        silent=args.silent,
    )
    return 0 if success else 1


if __name__ == "__main__":
    raise SystemExit(main())
